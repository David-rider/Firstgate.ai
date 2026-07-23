import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Sets the background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner padding for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_word_document(docx_path):
    doc = docx.Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Colors
    NAVY_DARK = RGBColor(14, 20, 34)      # #0E1422
    CYAN_ACCENT = RGBColor(0, 180, 216)    # #00B4D8
    SLATE_GRAY = RGBColor(100, 116, 139)  # #64748B
    TEXT_DARK = RGBColor(30, 41, 59)      # #1E293B

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("FirstGate.ai 官方网站规划与技术建设说明书")
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.name = 'Microsoft YaHei'
    r_title.font.color.rgb = NAVY_DARK

    # Document Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("Official Website Architecture, Feature Specs & Release Log v1.0")
    r_sub.font.size = Pt(12)
    r_sub.font.name = 'Arial'
    r_sub.font.color.rgb = CYAN_ACCENT

    # Meta Info Table
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False

    meta_data = [
        ("项目名称 / Project:", "FirstGate.ai 机构级 AI 算力交易市场与智能路由网关"),
        ("当前版本 / Version:", "v1.0.10 (Production Release)"),
        ("官方预览 / Live Preview:", "https://firstgate-ai.vercel.app/"),
        ("发布时间 / Date:", "2026年7月23日 (July 23, 2026)")
    ]

    for idx, (label, val) in enumerate(meta_data):
        row = table_meta.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "F8FAFC")
        set_cell_margins(cell_lbl, 80, 80, 120, 120)
        set_cell_margins(cell_val, 80, 80, 120, 120)

        p_lbl = cell_lbl.paragraphs[0]
        r_l = p_lbl.add_run(label)
        r_l.bold = True
        r_l.font.size = Pt(9.5)
        r_l.font.name = 'Microsoft YaHei'

        p_val = cell_val.paragraphs[0]
        r_v = p_val.add_run(val)
        r_v.font.size = Pt(9.5)
        r_v.font.name = 'Microsoft YaHei'

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for Headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.name = 'Microsoft YaHei'
        r.font.color.rgb = NAVY_DARK
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.name = 'Microsoft YaHei'
        r.font.color.rgb = CYAN_ACCENT
        return p

    def add_body_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.bold = True
            r_b.font.size = Pt(10.5)
            r_b.font.name = 'Microsoft YaHei'
            r_b.font.color.rgb = TEXT_DARK
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.name = 'Microsoft YaHei'
        r.font.color.rgb = TEXT_DARK
        return p

    # --- SECTION 1 ---
    add_heading_1("一、 项目定位与品牌核心倡导")
    add_body_p("FirstGate.ai 定位为机构级 AI 算力交易市场与多云智能路由网关。致力于解决大模型时代企业面临的异构 GPU 资源分散、算力采购成本高昂、多模型 API 调度复杂以及数据隐私安全合规等核心痛点。")
    add_body_p("品牌坚持「稳扎稳打、务实工程」原则，严禁编造任何夸大的虚拟数据（如虚假的数十万张 GPU 或百个数据中心），所有宣传与产品功能均建立在创始人 20 多年真实的 Enterprise IT、MSP 托管服务、路由交换组网、网络安全防御、金融合规（SOC2 / NY DFS）以及勒索病毒危机救援的硬核技术沉淀之上。")
    add_body_p("视觉美学采用了现代化暗黑高科技风格（Futuristic Cyber & Deep Navy），吸取了 PaleBlueDot 的极简精炼排版美学，兼具顶级科技震撼感与极佳的阅读舒适度。")

    # --- SECTION 2 ---
    add_heading_1("二、 前端技术栈与架构设计")
    add_heading_2("2.1 技术选型")
    add_body_p("全站基于无重型依赖的高效现代 Web 架构搭建，保证极佳的性能与亚毫秒级响应：")
    add_body_p("以语义化 HTML5 奠定结构，使用 Modular JS (ES Modules) 处理状态管理与动态交互。", "• 核心逻辑 (Core Engine): ")
    add_body_p("采用 Vanilla CSS 与 TailwindCSS CDN 组合，实现灵活的高科技赛博 UI 组件库。", "• 样式系统 (CSS System): ")
    add_body_p("采用 Vite 5.4.x 作为构建与热重载工具，编译产物压缩至 102kB，首屏秒开。", "• 打包与编译 (Build System): ")
    add_body_p("引入 Lucide Icons 矢量图标库与 Chart.js 实时网络时延与支出可视化图表。", "• 组件与可视化 (UI & Charts): ")
    add_body_p("基于 GitHub 代码仓库托管，自动触发 Vercel CI/CD 持续集成部署。", "• 自动化部署 (CI/CD Pipeline): ")

    add_heading_2("2.2 自研多语言国际化引擎 (i18n Engine)")
    add_body_p("网站内置了高可靠、无缝切换的自研动态轻量化 i18n 引擎（支持 English、简体中文、繁体中文 三语）：")
    add_body_p("通过 DOM 节点的 data-i18n 属性实现全自动文本替换，无需刷新页面。", "1. 零刷新即时切换：")
    add_body_p("支持普通文本、带 HTML 样式的富文本以及组件 Placeholder 的无缝多语言映射。", "2. 完整富文本映射：")
    add_body_p("经过全站零死角排查，所有页面文本、表格头、弹窗提示以及 JS 动态生成的日志行均 100% 接入字典。", "3. 100% 覆盖保证：")

    # --- SECTION 3 ---
    add_heading_1("三、 网站功能模块详细说明")
    
    table_modules = doc.add_table(rows=1, cols=3)
    table_modules.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_modules.autofit = False

    hdr_cells = table_modules.rows[0].cells
    hdr_titles = ["模块名称", "对应 Tab / ID", "功能描述与设计亮点"]
    hdr_widths = [Inches(1.5), Inches(1.5), Inches(3.5)]

    for idx, title in enumerate(hdr_titles):
        hdr_cells[idx].width = hdr_widths[idx]
        set_cell_background(hdr_cells[idx], "0E1422")
        set_cell_margins(hdr_cells[idx], 100, 100, 100, 100)
        p = hdr_cells[idx].paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.name = 'Microsoft YaHei'
        r.font.color.rgb = RGBColor(255, 255, 255)

    modules_info = [
        ("3D / 2D 智能大门封面页", "stargate-landing", "包含星门隧道动画、零信任网络隧道与硬件 TEE 飞地简介，作为震撼的品牌视觉入口。"),
        ("控制台概览主页", "tab-overview", "包含算力交易市场概览 Banner、四大架构支柱、真实客户端 RTT 测速图表及交互式路由仲裁模拟器。"),
        ("AI 算力交易 Spot 市场", "tab-marketplace", "展示实时 GPU (H100/H200/B200/L40S/A100) 竞价对比表及 SOTA 开源大模型一键部署目录。"),
        ("智能路由配置中心", "tab-router", "提供纽约极速低延迟、企业成本优化、物理隔离 VPC 三大预设策略及 Fallback 降级链可视化编排。"),
        ("华尔街安全与机密计算", "tab-security", "展示 SOC2/TEE/NY DFS/GDPR 四大合规标准、硬件 TEE 远程证明校验器、实时 PII 脱敏测试器及 ClickHouse 审计日志。"),
        ("关于我们与演进路线图", "tab-about", "展示创始人真实 20 年 IT/MSP/安全合规路线图（1990s至今四大时代）及单一底部终极预约行动卡片。"),
        ("开发文档与 SDK Quickstart", "tab-docs", "提供无缝替代 OpenAI SDK 的 Python/TS/cURL 示例代码（仅需修改 1 行 base_url）。"),
        ("节点验证与抓包证明 Modal", "modal-node-proof", "实时展示 Active Host、传输协议与测量 RTT，引导用户通过浏览器 F12 Network 验证 200 OK 真实性。"),
        ("API Key & Token 购买 Modal", "modal-portal-redirect", "全站统一的 API 密钥与 Token 额度购买对话框，基于 100% Viewport 垂直居中体系。")
    ]

    for mod_name, mod_id, mod_desc in modules_info:
        row_cells = table_modules.add_row().cells
        row_cells[0].width = hdr_widths[0]
        row_cells[1].width = hdr_widths[1]
        row_cells[2].width = hdr_widths[2]

        for i, val in enumerate([mod_name, mod_id, mod_desc]):
            set_cell_background(row_cells[i], "FAFAFA" if modules_info.index((mod_name, mod_id, mod_desc)) % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[i], 80, 80, 100, 100)
            p = row_cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.name = 'Microsoft YaHei'

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECTION 4 ---
    add_heading_1("四、 版本更新与优化履历 (Release Log)")
    
    releases = [
        ("v1.0.10", "精简合并 About Us 页面预约卡片", "删除了 About Us 顶部的重复预约按钮，保留底部统一的 Final Consultation CTA 卡片，使页面叙事逻辑流畅一致。"),
        ("v1.0.9", "统一全站弹窗 Viewport 100% 垂直居中", "重构 Modal 系统，将弹窗移至 <body> 根节点，引入 .global-modal-overlay，解决在不同滚动高度下弹窗高低不一的问题。"),
        ("v1.0.8", "限制 Sales Portal Banner 作用域", "将商业购买 Portal 横幅严格限制在 Overview 概览页面展示，避免在所有 Tab 重复堆叠，保持各子页面干练。"),
        ("v1.0.7", "修复 Logo 跳转与新增 Overview 独立 Tab", "纠正了左上角 Logo 点击误触发 switchTab('about') 的 BUG，新增 Overview 导航按钮，实现精准 1 对 1 导航映射。"),
        ("v1.0.6", "全站多语言 i18n 无死角审查与字典补全", "补充了 Security & TEE 页面缺失的 20 余项 zh-CN / zh-TW 翻译字典，实现全站多语言 100% 覆盖。"),
        ("v1.0.5", "顶栏全局吸顶固定与错位修复", "将顶部节点状态与语言切换条整合进 sticky top-0 容器，确保用户向下滚动时语言选择器始终置顶可见。"),
        ("v1.0.4", "整合精简 About Us 冗余卡片", "彻底删除了冗余的“Why FirstGate”重复卡片，将“年代演进”与“核心功底”融合成干练的 4 里程碑网格。"),
        ("v1.0.3", "统一导航按钮边框与恢复 Portal 位置", "为导航栏所有选项卡统一了 Rounded-XL 框线样式，并将 FirstGate Portal 还原至导航最左侧。"),
        ("v1.0.2", "务实创始人背景修正与全站虚假数据清除", "去除了关于服务器设计与热力学研发的描述，全面替换为创始人 20 年真实 Enterprise IT/MSP/网络安全救援功底；清除所有虚拟夸大数字。"),
        ("v1.0.1", "修复全站页面错位与 HTML 标签缺失", "排查并修复了 index.html 中未闭合的 <div> 标签，恢复了响应式 Flex 排版。"),
        ("v1.0.0", "官方网站初始版本上线", "搭建高科技赛博暗黑风格 UI 框架，完成多云算力交易、智能路由、TEE 安全、i18n 多语言引擎与 Vite 构建部署。")
    ]

    table_rel = doc.add_table(rows=1, cols=3)
    table_rel.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_rel.autofit = False

    r_hdr_cells = table_rel.rows[0].cells
    r_hdr_titles = ["版本号", "更新主题", "核心改进内容"]
    r_hdr_widths = [Inches(1.2), Inches(2.3), Inches(3.0)]

    for idx, title in enumerate(r_hdr_titles):
        r_hdr_cells[idx].width = r_hdr_widths[idx]
        set_cell_background(r_hdr_cells[idx], "0E1422")
        set_cell_margins(r_hdr_cells[idx], 100, 100, 100, 100)
        p = r_hdr_cells[idx].paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.name = 'Microsoft YaHei'
        r.font.color.rgb = RGBColor(255, 255, 255)

    for ver, topic, desc in releases:
        row_cells = table_rel.add_row().cells
        row_cells[0].width = r_hdr_widths[0]
        row_cells[1].width = r_hdr_widths[1]
        row_cells[2].width = r_hdr_widths[2]

        for i, val in enumerate([ver, topic, desc]):
            set_cell_background(row_cells[i], "FAFAFA" if releases.index((ver, topic, desc)) % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[i], 80, 80, 100, 100)
            p = row_cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.name = 'Microsoft YaHei'
            if i == 0:
                r.bold = True
                r.font.color.rgb = CYAN_ACCENT

    doc.save(docx_path)
    print(f"Word document saved successfully to: {docx_path}")

if __name__ == "__main__":
    output_dir = "d:\\study\\projects01\\Firstgate.ai\\docs"
    os.makedirs(output_dir, exist_ok=True)
    docx_file = os.path.join(output_dir, "FirstGate_AI_Release_Notes_v1.0.docx")
    create_word_document(docx_file)
